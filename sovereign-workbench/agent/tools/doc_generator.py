"""
Document Generator Tool for Karyalaya AI.

Responsibilities:
  - generate_approval_note()  — creates an AI draft (PENDING REVIEW). Agent-callable.
  - finalize_official_document() — writes the official .docx. Requires human_confirmed=True.
    Raises PermissionError otherwise. NOT callable by the agent loop (excluded from tool registry).
  - draft_payload_to_html()  — renders draft fields as HTML for inline UI preview.
  - create_approval_note_docx() — internal shared docx builder.
"""

import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from typing import Dict, Any, List

WORKSPACE_DIR = os.path.abspath(
    os.path.join(os.path.dirname(__file__), "..", "..", "workspace")
)
DELIVERABLES_DIR = os.path.join(WORKSPACE_DIR, "deliverables")


def create_approval_note_docx(
    title: str,
    ref_number: str,
    plant_unit: str,
    inspection_date: str,
    findings: List[str],
    sop_reference: str,
    recommendation: str,
    target_path: str,
    is_official: bool = False,
) -> str:
    """Builds and saves an approval note .docx at target_path. Returns the saved path."""
    doc = Document()

    for section in doc.sections:
        section.top_margin    = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin   = Inches(1)
        section.right_margin  = Inches(1)

    # Document type banner
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    banner = "OFFICIAL INSPECTION APPROVAL NOTE" if is_official \
        else "DRAFT INSPECTION APPROVAL NOTE (PENDING HUMAN REVIEW)"
    r_title = p_title.add_run(banner)
    r_title.bold = True
    r_title.font.size = Pt(15)
    r_title.font.color.rgb = RGBColor(0, 51, 102) if is_official else RGBColor(180, 100, 0)

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_sub = p_sub.add_run(f"KARYALAYA AI — {plant_unit.upper()}")
    r_sub.font.size = Pt(10)
    r_sub.font.italic = True

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # Metadata table
    table = doc.add_table(rows=4, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for idx, (label, val) in enumerate([
        ("Document Title:", title),
        ("Reference No.:", ref_number),
        ("Plant / Facility Unit:", plant_unit),
        ("Inspection Date:", inspection_date),
    ]):
        row = table.rows[idx]
        row.cells[0].text = label
        row.cells[1].text = val
        row.cells[0].paragraphs[0].runs[0].bold = True

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # Section 1: Findings
    h1 = doc.add_heading("1. Key Inspection Findings", level=1)
    h1.runs[0].font.color.rgb = RGBColor(0, 51, 102)
    for finding in findings:
        doc.add_paragraph(finding, style="List Bullet")

    # Section 2: SOP Grounding
    h2 = doc.add_heading("2. Grounding & SOP Reference", level=1)
    h2.runs[0].font.color.rgb = RGBColor(0, 51, 102)
    doc.add_paragraph(sop_reference)

    # Section 3: Recommendation
    h3 = doc.add_heading("3. Recommendation & Next Steps", level=1)
    h3.runs[0].font.color.rgb = RGBColor(0, 51, 102)
    doc.add_paragraph(recommendation)

    # Sign-off block
    doc.add_paragraph().paragraph_format.space_after = Pt(20)
    sign_text = (
        "Approved & Signed by:\n\n_______________________\nChief Inspection Officer / Unit Head"
        if is_official else
        "PENDING HUMAN REVIEW & SIGN-OFF\n\n_______________________\nAuthorized Approver Signature Required"
    )
    p_sign = doc.add_paragraph()
    p_sign.add_run(sign_text).bold = True

    doc.save(target_path)
    return target_path


def generate_approval_note(
    title: str,
    ref_number: str,
    plant_unit: str,
    inspection_date: str,
    findings: List[str],
    sop_reference: str,
    recommendation: str,
    filename: str = "Approval_Note.docx",
) -> Dict[str, Any]:
    """
    Creates an AI-drafted approval note in workspace/deliverables/ as a DRAFT.
    Returns a draft_payload dict that must be passed to finalize_official_document()
    after explicit human confirmation. This function is safe to call from the agent loop.
    """
    try:
        os.makedirs(DELIVERABLES_DIR, exist_ok=True)
        draft_filename = f"DRAFT_{filename}"
        draft_path = os.path.join(DELIVERABLES_DIR, draft_filename)

        create_approval_note_docx(
            title=title,
            ref_number=ref_number,
            plant_unit=plant_unit,
            inspection_date=inspection_date,
            findings=findings,
            sop_reference=sop_reference,
            recommendation=recommendation,
            target_path=draft_path,
            is_official=False,
        )

        draft_payload = {
            "title": title,
            "ref_number": ref_number,
            "plant_unit": plant_unit,
            "inspection_date": inspection_date,
            "findings": findings,
            "sop_reference": sop_reference,
            "recommendation": recommendation,
            "target_filename": filename,
        }

        return {
            "status": "draft_created",
            "requires_human_approval": True,
            "message": "AI Draft generated. Human sign-off required before finalizing.",
            "draft_file": draft_filename,
            "draft_payload": draft_payload,
        }
    except Exception as e:
        return {"status": "error", "error": str(e)}


def finalize_official_document(
    draft_payload: Dict[str, Any],
    *,
    human_confirmed: bool,
) -> Dict[str, Any]:
    """
    Writes the official .docx deliverable to workspace/deliverables/.

    ENFORCEMENT GATE: raises PermissionError if human_confirmed is not exactly True.
    This function is intentionally excluded from the agent tool registry and is only
    reachable via the /api/approve_draft HTTP endpoint, which supplies human_confirmed=True
    after an explicit user action in the UI.
    """
    if human_confirmed is not True:
        raise PermissionError(
            "Human confirmation is required to finalize an official document. "
            "This function must not be called without explicit human sign-off."
        )

    try:
        os.makedirs(DELIVERABLES_DIR, exist_ok=True)
        filename = draft_payload.get("target_filename", "Official_Approval_Note.docx")
        if not filename.endswith(".docx"):
            filename += ".docx"
        target_path = os.path.join(DELIVERABLES_DIR, filename)

        create_approval_note_docx(
            title=draft_payload.get("title", "Inspection Approval Note"),
            ref_number=draft_payload.get("ref_number", "REF-2026-001"),
            plant_unit=draft_payload.get("plant_unit", "Plant Unit"),
            inspection_date=draft_payload.get("inspection_date", "2026-08-27"),
            findings=draft_payload.get("findings", []),
            sop_reference=draft_payload.get("sop_reference", "Per SOP guidelines"),
            recommendation=draft_payload.get("recommendation", "Approved"),
            target_path=target_path,
            is_official=True,
        )

        return {
            "status": "success",
            "official_filename": filename,
            "file_path": target_path,
            "message": f"Human approval confirmed. Official deliverable '{filename}' created.",
        }
    except PermissionError:
        raise
    except Exception as e:
        return {"status": "error", "error": str(e)}


def draft_payload_to_html(draft_payload: Dict[str, Any], revision: int = 0) -> str:
    """
    Renders a draft_payload dict as structured HTML for inline browser preview.
    No file I/O — pure string rendering. Used by the /api/chat SSE stream.
    """
    revision_label = (
        f'<span class="revision-label">Revision {revision}</span>'
        if revision > 0
        else '<span class="revision-label">Initial Draft</span>'
    )
    findings = draft_payload.get("findings", [])
    findings_html = "".join(f"<li>{f}</li>" for f in findings) if findings else "<li>No findings recorded.</li>"

    return f"""<div class="doc-preview">
  <div class="doc-header">
    <div class="doc-status-bar">
      <span class="doc-draft-tag">AI DRAFT — PENDING HUMAN SIGN-OFF</span>
      {revision_label}
    </div>
    <div class="doc-title">{draft_payload.get("title", "Inspection Approval Note")}</div>
  </div>
  <table class="doc-meta-table">
    <tr><th>Reference No.</th><td>{draft_payload.get("ref_number", "—")}</td></tr>
    <tr><th>Plant / Unit</th><td>{draft_payload.get("plant_unit", "—")}</td></tr>
    <tr><th>Inspection Date</th><td>{draft_payload.get("inspection_date", "—")}</td></tr>
  </table>
  <div class="doc-section">
    <h3>Key Inspection Findings</h3>
    <ul>{findings_html}</ul>
  </div>
  <div class="doc-section">
    <h3>SOP Grounding &amp; Reference</h3>
    <p>{draft_payload.get("sop_reference", "—")}</p>
  </div>
  <div class="doc-section">
    <h3>Recommendation</h3>
    <p>{draft_payload.get("recommendation", "—")}</p>
  </div>
  <div class="doc-signoff">
    <p>PENDING HUMAN REVIEW &amp; SIGN-OFF</p>
    <div class="signoff-line">___________________________</div>
    <p class="signoff-role">Authorized Approver Signature Required</p>
  </div>
</div>"""
